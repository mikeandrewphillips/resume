#!/usr/bin/env python3
"""Generate a resume in Markdown, LaTeX, PDF, and Word.

resume.yaml is the complete master CV. A variant file in variants/ selects
which tagged elements to include and names the output.

Usage:
    python build.py [md|tex|pdf|docx|all] [variant]   (defaults: all full)

Outputs land in dist/. PDF requires xelatex on PATH (BasicTeX/MacTeX);
everything else is pure Python.
"""
from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

import yaml
from jinja2 import Environment, FileSystemLoader, StrictUndefined


class ToolchainError(RuntimeError):
    """Raised when an external tool (xelatex) is missing or fails."""


ROOT = Path(__file__).parent.resolve()
SRC = ROOT / "resume.yaml"
VARIANTS = ROOT / "variants"
TPL = ROOT / "templates"
DIST = ROOT / "dist"
CLS = ROOT / "latex" / "deedy-resume-openfont.cls"
FONTS = ROOT / "fonts"

# LaTeX special characters that must be escaped in user-supplied text.
_TEX_REPLACEMENTS = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def tex_escape(value: object) -> str:
    return "".join(_TEX_REPLACEMENTS.get(ch, ch) for ch in str(value))


# --------------------------------------------------------------------------
# Loading + tag filtering
# --------------------------------------------------------------------------
def _load_yaml(path: Path) -> dict:
    if not path.exists():
        sys.exit(f"error: {path} not found")
    with path.open(encoding="utf-8") as fh:
        return yaml.safe_load(fh) or {}


def load_cv() -> dict:
    return _load_yaml(SRC)


def load_variant(name: str) -> dict:
    return _load_yaml(VARIANTS / f"{name}.yaml")


def _selected(tags, selected) -> bool:
    """An element is kept when filtering is off, the element is core
    (untagged), or it carries at least one of the variant's tags."""
    if selected is None:        # full variant: no filtering
        return True
    if not tags:                # core element: always included
        return True
    return any(t in selected for t in tags)


def _entry(e):
    """Normalize a bullet/item: str -> (text, []); {text,tags} -> (text, tags)."""
    if isinstance(e, dict):
        return e.get("text", ""), e.get("tags") or []
    return e, []


def _filter_entries(entries, selected) -> list[str]:
    out = []
    for e in entries or []:
        text, tags = _entry(e)
        if _selected(tags, selected):
            out.append(text)
    return out


def _filter_groups(groups, selected) -> list[dict]:
    kept = []
    for g in groups or []:
        if not _selected(g.get("tags"), selected):
            continue
        items = _filter_entries(g.get("items"), selected)
        if items:
            kept.append({"name": g["name"], "items": items})
    return kept


def filter_cv(cv: dict, variant: dict) -> dict:
    selected = variant.get("tags")  # None -> keep everything

    education = [
        {"school": e["school"], "detail": e["detail"], "dates": e["dates"]}
        for e in cv.get("education", [])
        if _selected(e.get("tags"), selected)
    ]

    experience = []
    for company in cv.get("experience", []):
        if not _selected(company.get("tags"), selected):
            continue
        roles = []
        for role in company.get("roles", []):
            if not _selected(role.get("tags"), selected):
                continue
            bullets = _filter_entries(role.get("bullets"), selected)
            if bullets:
                roles.append(
                    {"title": role["title"], "dates": role["dates"], "bullets": bullets}
                )
        if roles:
            experience.append({"company": company["company"], "roles": roles})

    objective = variant.get("objective")
    if objective:
        objective = " ".join(objective.split())  # collapse YAML folded newlines

    return {
        "name": cv["name"],
        "contact": cv["contact"],
        "layout": variant.get("layout", "deedy"),
        "objective": objective,
        "education": education,
        "certifications": _filter_groups(cv.get("certifications"), selected),
        "technology": _filter_groups(cv.get("technology"), selected),
        "experience": experience,
        "publications": _filter_entries(cv.get("publications"), selected),
    }


# --------------------------------------------------------------------------
# Rendering
# --------------------------------------------------------------------------
def _markdown_env() -> Environment:
    return Environment(
        loader=FileSystemLoader(str(TPL)),
        trim_blocks=True,
        lstrip_blocks=True,
        undefined=StrictUndefined,
        keep_trailing_newline=True,
    )


def _latex_env() -> Environment:
    env = Environment(
        loader=FileSystemLoader(str(TPL)),
        block_start_string=r"\BLOCK{",
        block_end_string="}",
        variable_start_string=r"\VAR{",
        variable_end_string="}",
        comment_start_string=r"\#{",
        comment_end_string="}",
        trim_blocks=True,
        lstrip_blocks=True,
        autoescape=False,
        undefined=StrictUndefined,
        keep_trailing_newline=True,
    )
    env.filters["tex"] = tex_escape
    return env


def build_markdown(resume: dict, base: str) -> Path:
    out = DIST / f"{base}.md"
    rendered = _markdown_env().get_template("resume.md.j2").render(r=resume)
    out.write_text(rendered, encoding="utf-8")
    print(f"  wrote {out.relative_to(ROOT)}")
    return out


def build_latex(resume: dict, base: str) -> Path:
    out = DIST / f"{base}.tex"
    classic = resume.get("layout") == "classic"
    template = "cv.tex.j2" if classic else "resume.tex.j2"
    rendered = _latex_env().get_template(template).render(r=resume)
    out.write_text(rendered, encoding="utf-8")
    if not classic:  # Deedy needs its class file + bundled fonts staged
        shutil.copyfile(CLS, DIST / CLS.name)
        if FONTS.is_dir():
            shutil.copytree(FONTS, DIST / "fonts", dirs_exist_ok=True)
    print(f"  wrote {out.relative_to(ROOT)}")
    return out


def build_pdf(resume: dict, base: str) -> Path:
    if shutil.which("xelatex") is None:
        raise ToolchainError(
            "xelatex not found.\n"
            "  Install BasicTeX:  brew install --cask basictex\n"
            "  Then deps:         make tlmgr-deps"
        )
    build_latex(resume, base)
    # Two passes so textpos absolute positioning settles.
    for _ in range(2):
        proc = subprocess.run(
            ["xelatex", "-interaction=nonstopmode", "-halt-on-error", f"{base}.tex"],
            cwd=DIST,
            capture_output=True,
            text=True,
        )
    if proc.returncode != 0:
        log = DIST / f"{base}.log"
        tail = log.read_text(errors="replace").splitlines()[-25:] if log.exists() else []
        raise ToolchainError("xelatex failed:\n" + "\n".join(tail))
    for ext in (".aux", ".log", ".out"):
        (DIST / f"{base}{ext}").unlink(missing_ok=True)
    out = DIST / f"{base}.pdf"
    print(f"  wrote {out.relative_to(ROOT)}")
    return out


def build_docx(resume: dict, base: str) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor

    HEADING = RGBColor(0x6A, 0x6A, 0x6A)
    SUB = RGBColor(0x33, 0x33, 0x33)
    DATE = RGBColor(0x66, 0x66, 0x66)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(36)
        section.left_margin = section.right_margin = Pt(54)

    def hr(paragraph) -> None:
        p = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")
        borders.append(bottom)
        p.append(borders)

    def section_heading(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = HEADING

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = name.add_run(resume["name"])
    nr.font.size = Pt(28)
    nr.font.color.rgb = SUB

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = resume["contact"]
    cr = contact.add_run(f"{c['address']} | {c['phone']} | {c['email']}")
    cr.font.size = Pt(9)
    cr.font.color.rgb = DATE
    hr(contact)

    if resume.get("objective"):
        section_heading("Objective")
        op = doc.add_paragraph()
        op.paragraph_format.space_after = Pt(2)
        op.add_run(resume["objective"])

    section_heading("Experience")
    for company in resume["experience"]:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(8)
        cp.paragraph_format.space_after = Pt(0)
        cr = cp.add_run(company["company"])
        cr.bold = True
        cr.font.size = Pt(12)
        cr.font.color.rgb = SUB
        for role in company["roles"]:
            rp = doc.add_paragraph()
            rp.paragraph_format.space_after = Pt(2)
            rr = rp.add_run(f"{role['title']}  |  {role['dates']}")
            rr.italic = True
            rr.font.color.rgb = DATE
            for bullet in role["bullets"]:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)
                bp.add_run(bullet)

    if resume["education"]:
        section_heading("Education")
        for e in resume["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            er = p.add_run(e["school"])
            er.bold = True
            p.add_run(f" — {e['detail']} ")
            dr = p.add_run(f"({e['dates']})")
            dr.italic = True
            dr.font.color.rgb = DATE

    for title, key in (("Certifications", "certifications"), ("Technology", "technology")):
        if not resume[key]:
            continue
        section_heading(title)
        for group in resume[key]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(0)
            gr = p.add_run(group["name"])
            gr.bold = True
            gr.font.color.rgb = SUB
            for item in group["items"]:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(1)
                bp.add_run(item)

    if resume["publications"]:
        section_heading("Publications")
        for pub in resume["publications"]:
            pp = doc.add_paragraph(style="List Bullet")
            pp.paragraph_format.space_after = Pt(2)
            pp.add_run(pub)

    out = DIST / f"{base}.docx"
    doc.save(str(out))
    print(f"  wrote {out.relative_to(ROOT)}")
    return out


BUILDERS = {
    "md": build_markdown,
    "tex": build_latex,
    "pdf": build_pdf,
    "docx": build_docx,
}


def main(argv: list[str]) -> None:
    target = argv[0] if argv else "all"
    variant_name = argv[1] if len(argv) > 1 else "standard"
    if target not in {*BUILDERS, "all"}:
        sys.exit(f"error: unknown target {target!r}; choose md|tex|pdf|docx|all")

    DIST.mkdir(exist_ok=True)
    variant = load_variant(variant_name)
    resume = filter_cv(load_cv(), variant)
    base = variant.get("output", "resume")

    # pdf runs last so a missing LaTeX toolchain never blocks the others.
    targets = ["md", "tex", "docx", "pdf"] if target == "all" else [target]
    print(f"building [{variant_name} -> {base}]: {', '.join(targets)}")
    for t in targets:
        try:
            BUILDERS[t](resume, base)
        except ToolchainError as exc:
            if target == "all":
                print(f"  skipped {t}: {exc}")
            else:
                sys.exit(f"error: {exc}")
    print("done.")


if __name__ == "__main__":
    main(sys.argv[1:])
